"""
Eksport av håndbøker og risikovurdering til JSON, DOCX og PDF.

Markdown parses ÉN gang til en blokkliste, og samme blokkliste rendres til
både Word og PDF. Da kan ikke formatene komme ut av sync.
"""
import re
import json
import time
from pathlib import Path

BLAA = (31, 78, 121)   # overskriftsfarge, samme som Word-skjemaene
GRAA = (130, 130, 130)

_HMS_MAAL_KOLONNER = ["mål", "måltall", "frist", "ansvarlig"]


# ─── Markdown → blokker ──────────────────────────────────────────────────────

def _delt_tabellrad(linje: str) -> list[str]:
    return [c.strip() for c in linje.strip().strip("|").split("|")]


def _er_skillerad(linje: str) -> bool:
    """Rad av typen |---|---| som skiller header fra data."""
    return bool(re.fullmatch(r"\|?[\s:|-]+\|?", linje.strip())) and "-" in linje


def parse_markdown(md: str) -> list[dict]:
    """Del markdown i blokker: heading, avsnitt, punktliste, tabell, skillelinje."""
    blokker: list[dict] = []
    linjer = md.split("\n")
    i = 0
    while i < len(linjer):
        linje = linjer[i]
        strippet = linje.strip()

        if not strippet:
            i += 1
            continue

        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", strippet):
            blokker.append({"type": "hr"})
            i += 1
            continue

        overskrift = re.match(r"^(#{1,6})\s+(.*)$", strippet)
        if overskrift:
            blokker.append({
                "type": "heading",
                "nivaa": len(overskrift.group(1)),
                "tekst": overskrift.group(2).strip(),
            })
            i += 1
            continue

        # Tabell: minst headerrad + skillerad
        if strippet.startswith("|") and i + 1 < len(linjer) and _er_skillerad(linjer[i + 1]):
            headers = _delt_tabellrad(strippet)
            rader = []
            i += 2
            while i < len(linjer) and linjer[i].strip().startswith("|"):
                rader.append(_delt_tabellrad(linjer[i]))
                i += 1
            blokker.append({"type": "tabell", "headers": headers, "rader": rader})
            continue

        if re.match(r"^[-*+]\s+", strippet) or re.match(r"^\d+[.)]\s+", strippet):
            punkter = []
            nummerert = bool(re.match(r"^\d+[.)]\s+", strippet))
            while i < len(linjer):
                s = linjer[i].strip()
                m = re.match(r"^(?:[-*+]|\d+[.)])\s+(.*)$", s)
                if not m:
                    break
                punkter.append(m.group(1).strip())
                i += 1
            blokker.append({"type": "punktliste", "punkter": punkter, "nummerert": nummerert})
            continue

        # Avsnitt — samle sammenhengende linjer
        avsnitt = [strippet]
        i += 1
        while i < len(linjer):
            s = linjer[i].strip()
            if (not s or s.startswith("#") or s.startswith("|")
                    or re.match(r"^(?:[-*+]|\d+[.)])\s+", s)
                    or re.fullmatch(r"-{3,}|\*{3,}|_{3,}", s)):
                break
            avsnitt.append(s)
            i += 1
        blokker.append({"type": "avsnitt", "tekst": " ".join(avsnitt)})

    return blokker


def _inline_deler(tekst: str) -> list[tuple[str, bool]]:
    """Del tekst i (bit, fet) — håndterer **fet**."""
    deler = []
    for bit in re.split(r"(\*\*[^*]+\*\*)", tekst):
        if not bit:
            continue
        if bit.startswith("**") and bit.endswith("**") and len(bit) > 4:
            deler.append((bit[2:-2], True))
        else:
            deler.append((bit, False))
    return deler or [(tekst, False)]


def _til_reportlab_markup(tekst: str) -> str:
    """**fet** → <b>fet</b>, og escape XML-tegn som reportlab tolker."""
    trygg = tekst.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", trygg)


# ─── HMS-mål ─────────────────────────────────────────────────────────────────

def finn_hms_maal(md: str) -> list[dict]:
    """
    Hent HMS-måltabellen (Mål | Måltall | Frist | Ansvarlig) fra markdown.
    Returnerer tom liste hvis tabellen ikke finnes.
    """
    for blokk in parse_markdown(md):
        if blokk["type"] != "tabell":
            continue
        normaliserte = [h.lower().strip("* ") for h in blokk["headers"]]
        if not all(kol in normaliserte for kol in _HMS_MAAL_KOLONNER):
            continue
        idx = {kol: normaliserte.index(kol) for kol in _HMS_MAAL_KOLONNER}
        maal = []
        for rad in blokk["rader"]:
            if len(rad) < len(blokk["headers"]):
                continue
            maal.append({
                "maal":      rad[idx["mål"]].strip("* "),
                "maaltall":  rad[idx["måltall"]].strip("* "),
                "frist":     rad[idx["frist"]].strip("* "),
                "ansvarlig": rad[idx["ansvarlig"]].strip("* "),
            })
        return [m for m in maal if any(m.values())]
    return []


def hms_maal_feil(md: str, minimum: int = 3) -> list[str]:
    """
    Kvalitetsport for IK-forskriften § 5 andre ledd nr. 4: målene skal være konkrete og målbare.
    Krever tabell med minst `minimum` mål der måltall og frist inneholder tall.
    """
    maal = finn_hms_maal(md)
    if not maal:
        return ["Mangler HMS-måltabell med kolonnene Mål | Måltall | Frist | Ansvarlig"]
    feil = []
    if len(maal) < minimum:
        feil.append(f"Kun {len(maal)} HMS-mål — minimum {minimum} kreves")
    for m in maal:
        if not re.search(r"\d", m["maaltall"]):
            feil.append(f"HMS-målet «{m['maal']}» har ikke et tallfestet måltall")
        if not re.search(r"\d", m["frist"]):
            feil.append(f"HMS-målet «{m['maal']}» har ikke en datofestet frist")
        if not m["ansvarlig"]:
            feil.append(f"HMS-målet «{m['maal']}» mangler ansvarlig")
    return feil


# ─── IK-forskriften § 5: dokumentasjonskravene ───────────────────────────────

# Internkontrollforskriften § 5 andre ledd nr. 4–8 er de punktene som SKAL
# dokumenteres skriftlig. Hvert krav må ha et EGET kapittel — det er ikke nok at
# ordene finnes spredt i standardtekst. Derfor kreves treff både i kapittel-
# overskriften og i kapittelinnholdet. Hver gruppe er synonymer (minst én må treffe).
IK_DOKUMENTASJONSKRAV = [
    {
        "nr": 4,
        "navn": "Mål for helse, miljø og sikkerhet",
        "overskrift": ("hms-policy", "hms-mål", "policy", "innledning", "målsetting"),
        "innhold": [("mål",), ("frist", "innen", "ansvarlig")],
    },
    {
        "nr": 5,
        "navn": "Oversikt over organisasjon, ansvar, oppgaver og myndighet",
        "overskrift": ("ansvar", "organisering", "organisasjon", "roller"),
        "innhold": [("ansvar",), ("daglig leder", "verneombud")],
    },
    {
        "nr": 6,
        "navn": "Kartlegging av farer, risikovurdering, planer og tiltak",
        "overskrift": ("risiko", "kartlegging"),
        "innhold": [("risiko",), ("tiltak",)],
    },
    {
        "nr": 7,
        "navn": "Rutiner for å avdekke, rette opp og forebygge avvik",
        "overskrift": ("avvik",),
        "innhold": [("avvik",), ("melde", "meldeplikt", "rapporter", "rette opp", "korriger")],
    },
    {
        "nr": 8,
        "navn": "Systematisk overvåking og gjennomgang av internkontrollen",
        "overskrift": ("revisjon", "gjennomgang", "forbedring", "evaluering", "oppfølging"),
        "innhold": [("gjennomgå", "gjennomgang", "revisjon"), ("årlig", "hvert år", "én gang per år")],
    },
]


def _kapittelseksjoner(md: str) -> list[tuple[str, str]]:
    """Del dokumentet i (overskrift, kropp) per «## »-kapittel."""
    seksjoner: list[tuple[str, str]] = []
    overskrift, kropp = None, []
    for linje in md.split("\n"):
        if re.match(r"^##\s+\S", linje) and not linje.startswith("###"):
            if overskrift is not None:
                seksjoner.append((overskrift, "\n".join(kropp)))
            overskrift, kropp = linje.lstrip("#").strip(), []
        elif overskrift is not None:
            kropp.append(linje)
    if overskrift is not None:
        seksjoner.append((overskrift, "\n".join(kropp)))
    return seksjoner


def ik_dekning_feil(md: str) -> list[str]:
    """
    Kvalitetsport: hvert dokumentasjonskrav i IK-forskriften § 5 andre ledd nr. 4–8
    må ha et eget kapittel — treff kreves både i overskrift og innhold, slik at
    generell standardtekst ikke kan «dekke» et krav som mangler.
    """
    seksjoner = [(o.lower(), k.lower()) for o, k in _kapittelseksjoner(md)]
    if not seksjoner:
        return ["Dokumentet har ingen kapitler («## »-overskrifter) å kontrollere"]

    feil = []
    for krav in IK_DOKUMENTASJONSKRAV:
        kandidater = [k for o, k in seksjoner if any(h in o for h in krav["overskrift"])]
        if not kandidater:
            feil.append(
                f"IK-forskriften § 5 andre ledd nr. {krav['nr']} ({krav['navn']}): "
                f"mangler et kapittel om temaet"
            )
            continue
        if not any(
            all(any(ord_ in kropp for ord_ in gruppe) for gruppe in krav["innhold"])
            for kropp in kandidater
        ):
            feil.append(
                f"IK-forskriften § 5 andre ledd nr. {krav['nr']} ({krav['navn']}): "
                f"kapitlet finnes, men innholdet dekker ikke kravet"
            )
    return feil


# ─── Risikodata ──────────────────────────────────────────────────────────────

def risiko_rader(harvey_data: dict | None) -> list[dict]:
    """Bygg risikoradene fra Harveys analyse. Delt av Excel-, JSON-, Word- og PDF-eksport."""
    rader = []
    if not harvey_data:
        return rader
    for rf in harvey_data.get("risikofaktorer", []):
        rader.append({
            "prosess":     rf.get("faktor", ""),
            "risiko":      rf.get("faktor", ""),
            "hvem":        "Ansatte",
            "alvorlighet": rf.get("alvorlighet", "middels"),
            "tiltak":      rf.get("tiltak", ""),
        })
    for krav in harvey_data.get("bransjespesifikke_krav", []):
        rader.append({
            "prosess":     krav.get("krav", ""),
            "risiko":      krav.get("krav", ""),
            "hvem":        krav.get("gjelder_naar", krav.get("gjelder_når", "Ansatte")),
            "alvorlighet": "middels",
            "tiltak":      f"Iverksett iht. {krav.get('hjemmel', 'gjeldende regelverk')}",
        })
    return rader


# ─── JSON ────────────────────────────────────────────────────────────────────

def _dokumentmeta(tittel: str, dok_type: str) -> dict:
    return {
        "tittel": tittel,
        "type": dok_type,
        "versjon": "1.0",
        "dato": time.strftime("%Y-%m-%d"),
    }


def handbok_json(company_info: dict, tittel: str, dok_type: str,
                 kapitler: list[tuple[dict, str]], harvey_data: dict | None) -> dict:
    """Maskinlesbar håndbok — kapittelstruktur, HMS-mål og lovgrunnlag."""
    hd = harvey_data or {}
    samlet_md = "\n\n".join(tekst for _, tekst in kapitler)
    return {
        "bedrift": {
            "navn":           company_info.get("bedriftsnavn", ""),
            "organisasjonsnummer": company_info.get("orgnr", ""),
            "nace_kode":      company_info.get("nace_kode", ""),
            "bransje":        company_info.get("bransje", ""),
            "antall_ansatte": company_info.get("antall_ansatte"),
            "kontaktperson":  company_info.get("kontaktperson", ""),
        },
        "dokument": _dokumentmeta(tittel, dok_type),
        "hms_maal": finn_hms_maal(samlet_md) if dok_type == "hms" else [],
        "kapitler": [
            {
                "nummer":            kap.get("nummer"),
                "tittel":            kap.get("tittel", ""),
                "formaal":           kap.get("formaal", ""),
                "hjemler":           kap.get("hjemler", []),
                "stikkord":          kap.get("stikkord", []),
                "innhold_markdown":  tekst,
            }
            for kap, tekst in kapitler
        ],
        "lovgrunnlag": {
            "lover_alltid_gjeldende":  hd.get("lover_alltid_gjeldende", []),
            "bransjespesifikke_krav":  hd.get("bransjespesifikke_krav", []),
            "risikofaktorer":          hd.get("risikofaktorer", []),
        },
    }


def risiko_json(company_info: dict, harvey_data: dict | None) -> dict:
    neste = time.strftime("%Y-%m-%d", time.localtime(time.time() + 365 * 24 * 3600))
    meta = _dokumentmeta("Risikovurdering", "risikovurdering")
    meta["neste_revisjon"] = neste
    return {
        "bedrift": {
            "navn":      company_info.get("bedriftsnavn", ""),
            "nace_kode": company_info.get("nace_kode", ""),
            "bransje":   company_info.get("bransje", ""),
        },
        "dokument": meta,
        "skala": {
            "sannsynlighet": "1–5 (1 = svært lite sannsynlig, 5 = svært sannsynlig)",
            "konsekvens":    "1–5 (1 = ubetydelig, 5 = katastrofal)",
            "risikoscore":   "sannsynlighet × konsekvens",
        },
        "risikoer": [
            {
                "nr":             nr,
                "prosess":        r["prosess"],
                "risiko":         r["risiko"],
                "hvem_rammes":    r["hvem"],
                "alvorlighet":    r["alvorlighet"],
                "tiltak":         r["tiltak"],
                "sannsynlighet":  None,   # fylles ut av bedriften
                "konsekvens":     None,
                "risikoscore":    None,
                "ansvarlig":      "",
                "frist":          "",
                "status":         "Åpen",
            }
            for nr, r in enumerate(risiko_rader(harvey_data), 1)
        ],
    }


def skriv_json(data: dict, path: Path) -> Path:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


# ─── DOCX ────────────────────────────────────────────────────────────────────

def _docx_sidenummer(seksjon) -> None:
    """Legg «Side X av Y» i bunntekst via Word-feltkoder."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    p = seksjon.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def felt(instruksjon: str):
        start = OxmlElement("w:fldChar")
        start.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruksjon
        slutt = OxmlElement("w:fldChar")
        slutt.set(qn("w:fldCharType"), "end")
        run = p.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(*GRAA)
        for el in (start, instr, slutt):
            run._r.append(el)

    run = p.add_run("Side ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*GRAA)
    felt("PAGE")
    run = p.add_run(" av ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*GRAA)
    felt("NUMPAGES")


def til_docx(md: str, path: Path, bedriftsnavn: str, landskap: bool = False,
             sideskift_per_kapittel: bool = True) -> Path | None:
    """Render markdown til et Word-dokument klart til utskrift og signering."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
    except ImportError:
        return None

    doc = Document()
    for s in doc.sections:
        if landskap and s.page_width < s.page_height:
            s.orientation = WD_ORIENT.LANDSCAPE
            s.page_width, s.page_height = s.page_height, s.page_width
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.0 if landskap else 2.5)
        kolofon = s.header.paragraphs[0]
        kolofon.text = bedriftsnavn
        kolofon.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if kolofon.runs:
            kolofon.runs[0].font.size = Pt(8)
            kolofon.runs[0].font.color.rgb = RGBColor(*GRAA)
        _docx_sidenummer(s)

    normal = doc.styles["Normal"].font
    normal.name = "Calibri"
    normal.size = Pt(10.5)

    def skriv_avsnitt(p, tekst):
        for bit, fet in _inline_deler(tekst):
            run = p.add_run(bit)
            run.bold = fet

    for blokk in parse_markdown(md):
        if blokk["type"] == "heading":
            nivaa = min(max(blokk["nivaa"] - 1, 0), 4)
            h = doc.add_heading(blokk["tekst"], level=nivaa)
            for run in h.runs:
                run.font.color.rgb = RGBColor(*BLAA)
            if nivaa == 1 and sideskift_per_kapittel:
                h.paragraph_format.page_break_before = True

        elif blokk["type"] == "avsnitt":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            skriv_avsnitt(p, blokk["tekst"])

        elif blokk["type"] == "punktliste":
            stil = "List Number" if blokk["nummerert"] else "List Bullet"
            for punkt in blokk["punkter"]:
                p = doc.add_paragraph(style=stil)
                p.paragraph_format.space_after = Pt(2)
                skriv_avsnitt(p, punkt)

        elif blokk["type"] == "tabell":
            headers = blokk["headers"]
            t = doc.add_table(rows=1 + len(blokk["rader"]), cols=len(headers))
            t.style = "Table Grid"
            for i, h in enumerate(headers):
                celle = t.rows[0].cells[i]
                celle.text = ""
                run = celle.paragraphs[0].add_run(h.strip("* "))
                run.bold = True
                run.font.size = Pt(9)
            for r, rad in enumerate(blokk["rader"], 1):
                tom_celle = False
                for c in range(len(headers)):
                    celle = t.rows[r].cells[c]
                    celle.text = ""
                    verdi = rad[c] if c < len(rad) else ""
                    if not _fjern_markup(verdi):
                        tom_celle = True
                    for bit, fet in _inline_deler(verdi):
                        run = celle.paragraphs[0].add_run(bit)
                        run.bold = fet
                        run.font.size = Pt(9)
                if tom_celle:
                    # Skal fylles ut for hånd — gi raden skrivehøyde
                    t.rows[r].height = Cm(0.8)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

        elif blokk["type"] == "hr":
            continue  # sideskift håndteres av kapitteloverskriftene

    doc.save(str(path))
    return path


# ─── PDF ─────────────────────────────────────────────────────────────────────

def _kolonnebredder(headers: list[str], rader: list[list[str]], ncols: int,
                    tilgjengelig: float) -> list[float]:
    """
    Kolonnebredder som ikke deler ord midt i: hver kolonne får minst plass til
    sitt lengste ord, og resten fordeles etter hvor mye tekst kolonnen faktisk har.
    """
    from reportlab.pdfbase.pdfmetrics import stringWidth

    PADDING = 10.0
    minimum, vekt = [], []
    for c in range(ncols):
        celler = [headers[c]] + [r[c] for r in rader if c < len(r)]
        lengste_ord = 0.0
        for celle in celler:
            for ord_ in re.split(r"\s+", _fjern_markup(celle)):
                if ord_:
                    lengste_ord = max(lengste_ord, stringWidth(ord_, "Helvetica-Bold", 8.5))
        minimum.append(min(lengste_ord + PADDING, tilgjengelig * 0.35))
        vekt.append(max(sum(len(_fjern_markup(c2)) for c2 in celler) / max(len(celler), 1), 3))

    sum_min = sum(minimum)
    if sum_min >= tilgjengelig:
        # Får ikke plass uansett — skaler ned proporsjonalt
        return [tilgjengelig * m / sum_min for m in minimum]

    rest = tilgjengelig - sum_min
    sum_vekt = sum(vekt)
    return [m + rest * (v / sum_vekt) for m, v in zip(minimum, vekt)]


def _fjern_markup(tekst: str) -> str:
    return tekst.replace("**", "").strip("* ")


def til_pdf(md: str, path: Path, bedriftsnavn: str, dokumenttittel: str,
            landskap: bool = False, sideskift_per_kapittel: bool = True) -> Path | None:
    """Render markdown til PDF — samme blokker som Word-versjonen."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                        Table, TableStyle, PageBreak, KeepTogether)
    except ImportError:
        return None

    blaa = colors.Color(*[v / 255 for v in BLAA])
    graa = colors.Color(*[v / 255 for v in GRAA])

    ss = getSampleStyleSheet()
    stiler = {
        "brod": ParagraphStyle("brod", parent=ss["BodyText"], fontName="Helvetica",
                                fontSize=9.5, leading=13.5, spaceAfter=5),
        "h1":   ParagraphStyle("h1", parent=ss["Heading1"], fontName="Helvetica-Bold",
                                fontSize=17, leading=21, textColor=blaa, spaceAfter=10),
        "h2":   ParagraphStyle("h2", parent=ss["Heading2"], fontName="Helvetica-Bold",
                                fontSize=13, leading=17, textColor=blaa,
                                spaceBefore=10, spaceAfter=6),
        "h3":   ParagraphStyle("h3", parent=ss["Heading3"], fontName="Helvetica-Bold",
                                fontSize=11, leading=15, textColor=blaa,
                                spaceBefore=8, spaceAfter=4),
        "tittel": ParagraphStyle("tittel", parent=ss["Title"], fontName="Helvetica-Bold",
                                  fontSize=22, leading=27, textColor=blaa, alignment=TA_CENTER),
        "punkt": ParagraphStyle("punkt", parent=ss["BodyText"], fontName="Helvetica",
                                 fontSize=9.5, leading=13.5, leftIndent=14,
                                 bulletIndent=4, spaceAfter=3),
        "celle": ParagraphStyle("celle", parent=ss["BodyText"], fontName="Helvetica",
                                 fontSize=8.5, leading=11, spaceAfter=0),
        "celle_fet": ParagraphStyle("celle_fet", parent=ss["BodyText"], fontName="Helvetica-Bold",
                                     fontSize=8.5, leading=11, spaceAfter=0, textColor=colors.white),
    }

    sidestorrelse = landscape(A4) if landskap else A4

    def bunntekst(canvas, dok):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(graa)
        canvas.drawString(2.2 * cm, 1.3 * cm, bedriftsnavn)
        canvas.drawRightString(sidestorrelse[0] - 2.2 * cm, 1.3 * cm,
                               f"Side {canvas.getPageNumber()}")
        canvas.restoreState()

    dok = SimpleDocTemplate(
        str(path), pagesize=sidestorrelse,
        topMargin=2.0 * cm, bottomMargin=2.0 * cm,
        leftMargin=2.2 * cm, rightMargin=2.2 * cm,
        title=dokumenttittel, author=bedriftsnavn,
    )
    tilgjengelig = sidestorrelse[0] - 4.4 * cm

    flyt = []
    forste_kapittel = True
    for blokk in parse_markdown(md):
        if blokk["type"] == "heading":
            nivaa = blokk["nivaa"]
            tekst = _til_reportlab_markup(blokk["tekst"])
            if nivaa == 1:
                flyt.append(Paragraph(tekst, stiler["tittel"]))
                flyt.append(Spacer(1, 10))
            elif nivaa == 2:
                if not forste_kapittel and sideskift_per_kapittel:
                    flyt.append(PageBreak())
                forste_kapittel = False
                flyt.append(Paragraph(tekst, stiler["h1"]))
            else:
                flyt.append(Paragraph(tekst, stiler["h" + ("2" if nivaa == 3 else "3")]))

        elif blokk["type"] == "avsnitt":
            flyt.append(Paragraph(_til_reportlab_markup(blokk["tekst"]), stiler["brod"]))

        elif blokk["type"] == "punktliste":
            for n, punkt in enumerate(blokk["punkter"], 1):
                kule = f"{n}." if blokk["nummerert"] else "•"
                flyt.append(Paragraph(_til_reportlab_markup(punkt), stiler["punkt"], bulletText=kule))
            flyt.append(Spacer(1, 4))

        elif blokk["type"] == "tabell":
            headers = [h.strip("* ") for h in blokk["headers"]]
            ncols = len(headers)
            data = [[Paragraph(_til_reportlab_markup(h), stiler["celle_fet"]) for h in headers]]
            for rad in blokk["rader"]:
                data.append([
                    Paragraph(_til_reportlab_markup(rad[c] if c < len(rad) else ""), stiler["celle"])
                    for c in range(ncols)
                ])
            bredder = _kolonnebredder(headers, blokk["rader"], ncols, tilgjengelig)

            stil = [
                ("BACKGROUND", (0, 0), (-1, 0), blaa),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.Color(0.7, 0.7, 0.7)),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [colors.white, colors.Color(0.96, 0.97, 0.99)]),
            ]
            # Rader med tomme felt skal fylles ut for hånd — gi dem skrivehøyde
            for r, rad in enumerate(blokk["rader"], 1):
                if any(not _fjern_markup(rad[c] if c < len(rad) else "") for c in range(ncols)):
                    stil.append(("TOPPADDING", (0, r), (-1, r), 9))
                    stil.append(("BOTTOMPADDING", (0, r), (-1, r), 9))
            t = Table(data, colWidths=bredder, repeatRows=1)
            t.setStyle(TableStyle(stil))
            flyt.append(KeepTogether(t) if len(data) <= 8 else t)
            flyt.append(Spacer(1, 8))

        elif blokk["type"] == "hr":
            continue

    dok.build(flyt, onFirstPage=bunntekst, onLaterPages=bunntekst)
    return path


# ─── Risikovurdering som DOCX/PDF ────────────────────────────────────────────

def _risiko_markdown(company_info: dict, harvey_data: dict | None) -> str:
    """Bygg risikovurderingen som markdown, slik at docx/pdf-rendrerne kan gjenbrukes."""
    navn = company_info.get("bedriftsnavn", "Bedriften")
    bransje = company_info.get("bransje", "")
    nace = company_info.get("nace_kode", "")
    dato = time.strftime("%d.%m.%Y")
    neste = time.strftime("%d.%m.%Y", time.localtime(time.time() + 365 * 24 * 3600))

    linjer = [
        "# RISIKOVURDERING",
        f"## {navn}",
        f"**Bransje:** {bransje}" + (f" (NACE {nace})" if nace else "")
        + f"  |  **Utarbeidet:** {dato}  |  **Versjon:** 1.0  |  **Neste revisjon:** {neste}",
        "",
        "Sannsynlighet og konsekvens vurderes på skala 1–5. Risikoscore = "
        "sannsynlighet × konsekvens. Fyll ut kolonnene for sannsynlighet, konsekvens, "
        "ansvarlig og frist i fellesskap med verneombud og ansatte.",
        "",
        "| Nr | Farlig forhold / risiko | Hvem kan bli skadet | Alvorlighet | Forebyggende tiltak | S | K | Ansvarlig | Frist |",
        "|---|---|---|---|---|---|---|---|---|",
    ]
    for nr, r in enumerate(risiko_rader(harvey_data), 1):
        linjer.append(
            f"| {nr} | {r['risiko']} | {r['hvem']} | "
            f"{r['alvorlighet'].capitalize()} | {r['tiltak']} |  |  |  |  |"
        )
    linjer += [
        "",
        "## Signatur",
        "",
        "| Rolle | Navn | Dato | Signatur |",
        "|---|---|---|---|",
        "| Daglig leder |  |  |  |",
        "| Verneombud |  |  |  |",
    ]
    return "\n".join(linjer)


def skriv_risikovurdering(company_info: dict, harvey_data: dict | None,
                          output_dir: Path, safe_navn: str) -> list[Path]:
    """Risikovurdering som JSON, DOCX og PDF (Excel lages fortsatt av pipeline)."""
    navn = company_info.get("bedriftsnavn", "Bedriften")
    md = _risiko_markdown(company_info, harvey_data)
    filer = []

    filer.append(skriv_json(risiko_json(company_info, harvey_data),
                            output_dir / f"{safe_navn}_Risikovurdering.json"))
    for path in (
        til_docx(md, output_dir / f"{safe_navn}_Risikovurdering.docx", navn,
                 landskap=True, sideskift_per_kapittel=False),
        til_pdf(md, output_dir / f"{safe_navn}_Risikovurdering.pdf", navn,
                f"Risikovurdering — {navn}", landskap=True, sideskift_per_kapittel=False),
    ):
        if path:
            filer.append(path)
    return filer


def _arlig_revisjon_markdown(company_info: dict, harvey_data: dict | None) -> str:
    """
    Årlig gjennomgang av internkontrollen — IK-forskriften § 5 andre ledd nr. 8.
    Dette er dokumentet Arbeidstilsynet ber om for å se at systemet faktisk brukes.
    """
    navn = company_info.get("bedriftsnavn", "Bedriften")
    aar = time.strftime("%Y")
    linjer = [
        "# ÅRLIG GJENNOMGANG AV HMS-SYSTEMET",
        f"## {navn}",
        f"**Gjennomgangsår:** {aar}  |  **Hjemmel:** IK-forskriften § 5 andre ledd nr. 8",
        "",
        "Internkontrollen skal gjennomgås systematisk minst én gang per år for å bekrefte "
        "at den fungerer som forutsatt. Fyll ut skjemaet i fellesskap med verneombud, "
        "og arkiver det som dokumentasjon.",
        "",
        "## 1. Deltakere og dato",
        "",
        "| Rolle | Navn | Til stede |",
        "|---|---|---|",
        "| Daglig leder |  |  |",
        "| Verneombud |  |  |",
        "| Øvrige deltakere |  |  |",
        "",
        "| Dato for gjennomgang | Forrige gjennomgang |",
        "|---|---|",
        "|  |  |",
        "",
        "## 2. Dokumentasjonskravene i IK-forskriften § 5",
        "",
        "Kontroller at hvert krav er oppfylt, oppdatert og tilgjengelig for de ansatte.",
        "",
        "| Nr | Krav | Oppfylt (ja/nei) | Sist oppdatert | Kommentar |",
        "|---|---|---|---|---|",
    ]
    for krav in IK_DOKUMENTASJONSKRAV:
        linjer.append(f"| {krav['nr']} | {krav['navn']} |  |  |  |")
    linjer += [
        "",
        "## 3. HMS-målene fra i fjor",
        "",
        "| Mål | Måltall | Oppnådd? | Tiltak videre |",
        "|---|---|---|---|",
        "|  |  |  |  |",
        "|  |  |  |  |",
        "|  |  |  |  |",
        "",
        "## 4. Avvik og hendelser siste år",
        "",
        "| Antall meldte avvik | Antall lukket | Alvorlige hendelser meldt Arbeidstilsynet |",
        "|---|---|---|",
        "|  |  |  |",
        "",
        "**Gjentakende avvik og hva vi gjør med dem:**",
        "",
        "| Gjentakende avvik | Årsak | Tiltak | Ansvarlig |",
        "|---|---|---|---|",
        "|  |  |  |  |",
        "|  |  |  |  |",
        "",
        "## 5. Risikovurdering og vernerunder",
        "",
        "| Kontrollpunkt | Status | Dato | Kommentar |",
        "|---|---|---|---|",
        "| Risikovurderingen er oppdatert |  |  |  |",
        "| Vernerunder gjennomført som planlagt |  |  |  |",
        "| Tiltak fra handlingsplanen er gjennomført |  |  |  |",
        "| Nye risikoforhold er vurdert |  |  |  |",
    ]
    risikoer = risiko_rader(harvey_data)
    if risikoer:
        linjer += [
            "",
            "**Bransjespesifikke risikoforhold som skal vurderes særskilt:**",
            "",
        ]
        linjer += [f"- {r['risiko']} — {r['tiltak']}" for r in risikoer[:10]]
    linjer += [
        "",
        "## 6. Opplæring og medvirkning",
        "",
        "| Kontrollpunkt | Status | Kommentar |",
        "|---|---|---|",
        "| HMS-opplæring for daglig leder er gjennomført (AML § 3-5) |  |  |",
        "| Verneombudet har fått opplæring (AML § 6-5) |  |  |",
        "| Nyansatte har fått HMS-introduksjon |  |  |",
        "| De ansatte er informert om endringer i rutinene |  |  |",
        "",
        "## 7. Konklusjon",
        "",
        "**Fungerer internkontrollen som forutsatt?**   Ja ☐   Delvis ☐   Nei ☐",
        "",
        "**Nye HMS-mål og tiltak for neste år:**",
        "",
        "| Mål | Måltall | Frist | Ansvarlig |",
        "|---|---|---|---|",
        "|  |  |  |  |",
        "|  |  |  |  |",
        "|  |  |  |  |",
        "",
        "## 8. Signatur",
        "",
        "| Rolle | Navn | Dato | Signatur |",
        "|---|---|---|---|",
        "| Daglig leder |  |  |  |",
        "| Verneombud |  |  |  |",
    ]
    return "\n".join(linjer)


def skriv_arlig_revisjon(company_info: dict, harvey_data: dict | None,
                         output_dir: Path, safe_navn: str) -> list[Path]:
    """Skjema for den årlige gjennomgangen av internkontrollen (DOCX + PDF)."""
    navn = company_info.get("bedriftsnavn", "Bedriften")
    md = _arlig_revisjon_markdown(company_info, harvey_data)
    filer = []
    for path in (
        til_docx(md, output_dir / f"{safe_navn}_Årlig_HMS_revisjon.docx", navn,
                 sideskift_per_kapittel=False),
        til_pdf(md, output_dir / f"{safe_navn}_Årlig_HMS_revisjon.pdf", navn,
                f"Årlig HMS-revisjon — {navn}", sideskift_per_kapittel=False),
    ):
        if path:
            filer.append(path)
    return filer


def skriv_handbok(md: str, company_info: dict, tittel: str, dok_type: str,
                  kapitler: list[tuple[dict, str]], harvey_data: dict | None,
                  output_dir: Path, filbasis: str) -> list[Path]:
    """Håndbok som JSON, DOCX og PDF ved siden av markdown-versjonen."""
    navn = company_info.get("bedriftsnavn", "Bedriften")
    filer = [skriv_json(
        handbok_json(company_info, tittel, dok_type, kapitler, harvey_data),
        output_dir / f"{filbasis}.json",
    )]
    for path in (
        til_docx(md, output_dir / f"{filbasis}.docx", navn),
        til_pdf(md, output_dir / f"{filbasis}.pdf", navn, f"{tittel} — {navn}"),
    ):
        if path:
            filer.append(path)
    return filer
